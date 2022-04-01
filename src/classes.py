from pytube import YouTube
from moviepy.editor import VideoFileClip, AudioFileClip
#from moviepy import AudioFileClip
from docx import Document
import speech_recognition as sr



class Link:
    """Idea es que se guardé un link, y se descargue ya sea el audio o video de este link"""
    
    
    def __init__(self, url : str = "") -> None:
        self.url = url
        
    
    def descargar_audio(self):
        """Del respectivo link, va a descargar solamente el audio en un solo archivo"""
        
        yt = YouTube(self.url)
        print("El nombre del video es " + yt.title)
        print(yt.streams.filter(only_audio=True))
        tag = int(input("Ingrese el tag que quiere utilizar: "))
        stream = yt.streams.get_by_itag(tag)
        stream.download()
        print("El audio del video ha sido descargado de forma exitosa\n")
        
    
    def extraccion_formato_video_a_audio_mp3(self):
        """Sirve para extraer audio por ejemplo de un archivo .mp4 y lo escribe uno .mp3"""
        
        nombre_clip = input("Ingresa el nombre del archivo, ejemplo: Cancion1.mp4")
        my_clip = VideoFileClip(r"{nombre}".format(nombre=nombre_clip))
        my_clip.audio.write_audiofile("Audio.mp3")
        print("El audio ha sido separado de su video de forma exitosa\n")
        
    
    def conversion_formato_video_a_audio_wav(self):
        """Extrae el audio de un video y lo escribe en un formato .wav"""
        
        clip = input("Ingresa el nombre del archivo completo, ejemplo: Video1.mp4: ")
        video = VideoFileClip(clip)
        audio = video.audio
        nombre_archivo = input("Ingrese como quiere escribir el audio y su extension, Ej: Audio.wav: ")
        audio.write_audiofile(nombre_archivo)
        print("Se ha extraido el audio de su video de forma exitosa")


class Recognition():
    
    """Pones el nombre del primer archivo y el último. Ejemplo: Si hay Audio1.wav, Audio2.wav y Audio3.wav, pones Recognition(1, 3, Audio).
    El audio tiene que estar en formato .wav
    """
    
    def __init__(self, inicio: int, fin: int, nombre_documento: str = "") -> None:
        self.inicio = inicio
        self.fin = fin
        self.nombre = nombre_documento
        
        
    def reconocimiento(self):
        
        """Convierte todos los archivos de audio que tienen el mismo nombre, pero distinto número, muchos archivos word con el respectivo texto del audio"""
        
        r = sr.Recognizer()
        print("Recuerda que tienes que tener los documentos con el mismo nombre principal y distinto número. Ejemplo: Audio1.wav, Audio2.wav ...")
        
        for i in range(self.inicio, self.fin + 1):
            with sr.AudioFile(f"{self.nombre}{i}.wav") as source:
                audio_text = r.listen(source)
                
                try:
                    text = r.recognize_google(audio_text, language="es-CL")
                    print("Convirtiendo a texto")
                    document = Document()
                    document.add_paragraph(text)
                    document.save(f"Transcripcion{i}.docx")
                    print("La transcripción ha sido realizada con éxito")
                    print("")
                except:
                    print("Intenta de nuevo")
                    
                    
class ArchivoAudio:
    
    "Tiene que ser un archivo de audio junto con su extensión: Ejemplo Cancion.wav"
    
    def __init__(self, nombre_archivo):
        self.audio = nombre_archivo
        
    def division_partes_iguales(self):
        my_audio = AudioFileClip(self.audio)
        print(f"La duración del audio en segundos es de : {my_audio.duration} segundos")
        print(f"Puedes dividir el audio en {my_audio.duration / 180} partes de 3 minutos cada una")
        cantidad_final = int(my_audio.duration / 180)
        t_1 = 0
        t_2 = 180
        
        for i in range(1, cantidad_final + 1):
            clip_1 = my_audio.subclip(t_1, t_2)
            clip_1.write_audiofile(f"Parte {i}.wav")
            t_1 += 180
            t_2 += 180
            
            i += 1
        
        print(t_1)
        print(my_audio.duration)
        clip = my_audio.subclip(t_1, my_audio.duration/180 - 2)
        clip.write_audiofile(f"Parte {cantidad_final + 1}.wav")
        
        print(f"Se logró dividir el audio en partes iguales de forma exitosa\n")
        
        
class Texto:
    """Ingresa el nombre del documento, tiene que ser formato word, sin el .docx
    Se puede obtener el texto de un word, iterar sobre ellos y juntar toda la información
    """
    def __init__(self, nombre_documento, inicio, fin):
        self.filename = nombre_documento
        self.inicio = inicio
        self.fin = fin 
        self.fullText = []
        
    
    def getText(self, filename):
        """Obtiene el texto de un archivo word, creo"""
        
        
        doc = Document(filename)
        
        for para in doc.paragraphs:
            self.fullText.append(para.text)
        
        
        return '\n'.join(self.fullText)
    
    
    def loop_documentos(self):
        """Añade toda la información a la variable self.fullText, de todos los word"""
        try:
            for i in range(self.inicio, self.fin + 1):
                self.getText(f"{self.filename}{i}.docx")
        
        except:
            
            pass
        
        finally:
            print("El loop fue completado")
            
    
    def juntar_textos_en_uno(self):
        """Añade toda la información de self.fullText a una nueva variable documento"""
        
        documento = Document()
        
        for i in self.fullText:
            documento.add_paragraph(i)
            
        documento.save("Textos_en_uno.docx")
        print("Tus documentos han sido juntados con éxito")
        
    def juntar_words_en_uno(self):
        self.loop_documentos()
        print("Todavía no se juntan los documentos")
        self.juntar_textos_en_uno()